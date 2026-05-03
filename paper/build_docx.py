"""Convert paper/methods-note-draft.md → paper/methods-note-draft.docx.

Synthēsis house style: A4, 1.5 line spacing, 11-pt Calibri body,
12-pt Times New Roman headings. Figure embedded inline. Vancouver
references rendered as a numbered list.
"""
from __future__ import annotations

import re
import sys
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Mm, Pt, Inches

_FRONTMATTER_RX = re.compile(r"\A---\n.*?\n---\n", re.DOTALL)
_CAPTION_RX = re.compile(
    r"<!--\s*figure-caption-begin\s*-->(.*?)<!--\s*figure-caption-end\s*-->",
    re.DOTALL | re.IGNORECASE,
)
_HTML_COMMENT_RX = re.compile(r"<!--.*?-->", re.DOTALL)
_BOLD_RX = re.compile(r"\*\*(.+?)\*\*")
_ITAL_RX = re.compile(r"\*(.+?)\*")


def _set_a4(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Mm(25)
    section.right_margin = Mm(25)
    section.top_margin = Mm(25)
    section.bottom_margin = Mm(25)


def _set_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    for h in ("Heading 1", "Heading 2"):
        s = doc.styles[h]
        s.font.name = "Times New Roman"
        s.font.size = Pt(12)


def _add_runs(p, text: str) -> None:
    """Add a paragraph with **bold** and *italic* spans honoured."""
    i = 0
    for m in _BOLD_RX.finditer(text):
        if m.start() > i:
            p.add_run(text[i:m.start()])
        run = p.add_run(m.group(1))
        run.bold = True
        i = m.end()
    rest = text[i:]
    j = 0
    for m in _ITAL_RX.finditer(rest):
        if m.start() > j:
            p.add_run(rest[j:m.start()])
        run = p.add_run(m.group(1))
        run.italic = True
        j = m.end()
    if j < len(rest):
        p.add_run(rest[j:])


def _add_paragraph(doc: Document, text: str, *, style: Optional[str] = None) -> None:
    p = doc.add_paragraph(style=style)
    _add_runs(p, text)


def export(md_path: Path, out_path: Path, figure_path: Optional[Path]) -> None:
    md = md_path.read_text(encoding="utf-8")
    md = _FRONTMATTER_RX.sub("", md, count=1)

    cap_m = _CAPTION_RX.search(md)
    caption = cap_m.group(1).strip() if cap_m else ""
    md_no_caption = _CAPTION_RX.sub("", md)
    md_no_caption = _HTML_COMMENT_RX.sub("", md_no_caption)

    doc = Document()
    _set_a4(doc)
    _set_styles(doc)

    body, _, refs = md_no_caption.partition("## References")
    figure_inserted = False

    for raw in body.splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        else:
            if (not figure_inserted
                    and figure_path is not None and figure_path.exists()
                    and line.lstrip().startswith("**Data availability")):
                doc.add_picture(str(figure_path), width=Inches(6.0))
                if caption:
                    cp = doc.add_paragraph()
                    _add_runs(cp, caption)
                figure_inserted = True
            _add_paragraph(doc, line)

    if not figure_inserted and figure_path is not None and figure_path.exists():
        doc.add_picture(str(figure_path), width=Inches(6.0))
        if caption:
            _add_paragraph(doc, caption)

    if refs.strip():
        doc.add_heading("References", level=2)
        for line in refs.splitlines():
            line = line.strip()
            m = re.match(r"^\d+\.\s+(.+)$", line)
            if m:
                _add_paragraph(doc, m.group(1))

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def main() -> int:
    md = Path("paper/methods-note-draft.md")
    out = Path("paper/methods-note-draft.docx")
    fig = Path("figures/figure-d.png")
    export(md, out, fig if fig.exists() else None)
    print(f"docx OK: {out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
